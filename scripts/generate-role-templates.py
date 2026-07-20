#!/usr/bin/env python3
"""Generate role-owned DOCX templates from the reviewed Vietnamese prompt pack.

The prompt pack remains the content source. This builder deliberately keeps all
unknown project facts as [ĐIỀN...] fields, applies one Power Home document
system, scrubs OOXML metadata, and emits the matching in-browser HTML preview.

Run one role at a time so the generation log and review scope remain explicit:

  python scripts/generate-role-templates.py --prompt-file <path> --role sa
"""

from __future__ import annotations

import argparse
import importlib.util
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_ROOT = ROOT / "assets" / "templates"
PREVIEW_ROOT = OUTPUT_ROOT / "previews"

POWER_RED = "DA2128"
POWER_DARK_RED = "A71920"
POWER_YELLOW = "F8D431"
INK = "25324A"
MUTED = "667085"
LIGHT_RED = "FDEDEE"
LIGHT_GRAY = "F4F6F8"
GRID = "CBD2D9"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9026  # A4 portrait, 1 inch left/right margins.
TABLE_INDENT_DXA = 120


@dataclass(frozen=True)
class PromptRecord:
    code: str
    heading: str
    pic: str
    inputs: tuple[str, ...]
    sections: tuple[str, ...]
    principles: tuple[str, ...]
    output: str


ROLE_SPECS = {
    "sa": {
        "label": "SOLUTION ARCHITECT",
        "codes": ("4.3", "4.4"),
        "files": {
            "4.3": "api-specification.docx",
            "4.4": "non-functional-requirements-specification.docx",
        },
    },
    "ux": {
        "label": "UX/UI DESIGNER",
        "codes": ("5.1", "5.2", "5.3", "5.4", "5.5"),
        "files": {
            "5.1": "user-persona-customer-journey-map.docx",
            "5.2": "information-architecture-wireframe-brief.docx",
            "5.3": "design-system-ui-style-guide.docx",
            "5.4": "usability-test-plan-report.docx",
            "5.5": "design-handoff-specification.docx",
        },
    },
    "qc": {
        "label": "QC TESTING",
        "codes": ("6.1", "6.2", "6.3", "6.4", "6.5"),
        "files": {
            "6.1": "test-strategy-test-plan.docx",
            "6.2": "test-case-specification.docx",
            "6.3": "defect-report.docx",
            "6.4": "test-summary-report.docx",
            "6.5": "uat-plan.docx",
        },
    },
    "sec": {
        "label": "SECURITY & COMPLIANCE",
        "codes": ("7.1", "7.2", "7.3", "7.4", "7.5"),
        "files": {
            "7.1": "threat-model-stride.docx",
            "7.2": "security-requirements-checklist.docx",
            "7.3": "security-risk-assessment.docx",
            "7.4": "incident-response-plan.docx",
            "7.5": "compliance-gap-checklist.docx",
        },
    },
    "ops": {
        "label": "OPERATIONS & SRE",
        "codes": ("8.1", "8.2", "8.3", "8.4"),
        "files": {
            "8.1": "sli-slo-sla-document.docx",
            "8.2": "operational-runbook.docx",
            "8.3": "blameless-postmortem.docx",
            "8.4": "disaster-recovery-plan.docx",
        },
    },
    "deployment": {
        "label": "DEPLOYMENT / DEVOPS",
        "codes": ("10.1", "10.2", "10.3", "10.4"),
        "files": {
            "10.1": "deployment-runbook-rollback-plan.docx",
            "10.2": "release-go-no-go-checklist.docx",
            "10.3": "ci-cd-pipeline-documentation.docx",
            "10.4": "environment-configuration-document.docx",
        },
    },
    "pmo": {
        "label": "PMO GOVERNANCE & PORTFOLIO",
        "codes": ("11.1", "11.2", "11.3", "11.4"),
        "files": {
            "11.1": "pmo-charter-governance-framework.docx",
            "11.2": "stage-gate-review-template.docx",
            "11.3": "portfolio-status-report.docx",
            "11.4": "benefits-realization-plan.docx",
        },
    },
}


# Each worksheet is genuine row/column data requested by its source prompt.
# Width weights are normalized to the exact A4 content width at render time.
WORKSHEETS = {
    "4.3": [
        ("Danh mục API", ("API ID", "Capability", "Consumer", "Owner", "Status"), (12, 31, 22, 20, 15)),
        ("Đặc tả endpoint", ("Method / Path", "Request", "Response", "Auth", "Error / SLA"), (22, 22, 22, 14, 20)),
    ],
    "4.4": [
        ("NFR Register", ("NFR ID", "Quality attribute", "Metric / Target", "Verification", "Owner"), (12, 23, 28, 23, 14)),
        ("Quality Attribute Scenario", ("Source", "Stimulus", "Environment", "Response", "Measure"), (18, 20, 18, 24, 20)),
    ],
    "5.1": [
        ("Persona profile", ("Persona", "Goal", "Pain point", "Behaviour", "Evidence"), (16, 21, 24, 20, 19)),
        ("Customer journey", ("Stage", "Action", "Touchpoint", "Emotion", "Opportunity"), (16, 22, 21, 16, 25)),
    ],
    "5.2": [
        ("Content inventory", ("Content / Screen", "User need", "Parent", "Priority", "Owner"), (24, 28, 18, 15, 15)),
        ("Wireframe register", ("ID", "Screen", "Purpose", "State / Variant", "Review status"), (10, 22, 30, 23, 15)),
    ],
    "5.3": [
        ("Design tokens", ("Token", "Value", "Usage", "Accessibility check"), (22, 18, 35, 25)),
        ("Component inventory", ("Component", "States", "Behaviour", "Owner", "Version"), (21, 20, 31, 16, 12)),
    ],
    "5.4": [
        ("Test scenarios", ("Task ID", "Scenario / Goal", "Success criteria", "Metric", "Moderator"), (12, 30, 27, 16, 15)),
        ("Findings", ("Finding", "Evidence", "Severity", "Recommendation", "Owner"), (24, 25, 13, 25, 13)),
    ],
    "5.5": [
        ("Handoff register", ("Screen / Flow", "Design link", "States", "Acceptance notes", "Dev owner"), (22, 20, 18, 27, 13)),
        ("Design QA checklist", ("Check", "Expected", "Evidence", "Result", "Owner"), (25, 28, 22, 12, 13)),
    ],
    "6.1": [
        ("Test scope", ("Feature / Risk", "Test level", "Test type", "Priority", "Owner"), (28, 18, 22, 14, 18)),
        ("Entry / Exit criteria", ("Gate", "Criterion", "Evidence", "Decision owner"), (16, 38, 28, 18)),
    ],
    "6.2": [
        ("Test cases", ("TC ID", "Precondition", "Steps", "Expected result", "Priority"), (12, 22, 29, 25, 12)),
        ("Traceability", ("TC ID", "Requirement / AC", "Test data", "Result", "Defect ID"), (15, 32, 23, 15, 15)),
    ],
    "6.3": [
        ("Defect record", ("Field", "Value / Evidence"), (25, 75)),
        ("Defect lifecycle", ("Status", "Owner", "Entry condition", "Exit condition"), (18, 18, 32, 32)),
    ],
    "6.4": [
        ("Execution summary", ("Test type", "Planned", "Executed", "Passed", "Failed"), (28, 18, 18, 18, 18)),
        ("Open defects", ("Severity", "Count", "Risk", "Disposition", "Owner"), (18, 13, 29, 24, 16)),
    ],
    "6.5": [
        ("UAT scenarios", ("UAT ID", "Business scenario", "Expected outcome", "Business owner", "Result"), (12, 31, 27, 18, 12)),
        ("UAT sign-off", ("Decision", "Condition / Exception", "Owner", "Due date"), (18, 44, 22, 16)),
    ],
    "7.1": [
        ("STRIDE register", ("Asset / Flow", "Threat", "STRIDE", "Risk", "Mitigation"), (23, 24, 15, 14, 24)),
        ("Trust boundaries", ("Boundary", "Data crossing", "Control", "Owner"), (24, 30, 30, 16)),
    ],
    "7.2": [
        ("Security requirements", ("SEC ID", "Requirement", "Verification", "Evidence", "Status"), (12, 34, 23, 19, 12)),
        ("Control checklist", ("Domain", "Control", "Mandatory", "Owner", "Result"), (21, 32, 15, 17, 15)),
    ],
    "7.3": [
        ("Security risk register", ("Risk ID", "Scenario", "Likelihood", "Impact", "Treatment"), (12, 34, 16, 14, 24)),
        ("Treatment plan", ("Action", "Control owner", "Due", "Residual risk", "Status"), (32, 18, 14, 22, 14)),
    ],
    "7.4": [
        ("Incident roles", ("Role", "Responsibility", "Primary", "Backup"), (20, 42, 19, 19)),
        ("Response playbook", ("Phase", "Trigger", "Actions", "Evidence", "Owner"), (15, 20, 30, 20, 15)),
    ],
    "7.5": [
        ("Compliance gaps", ("Control / Clause", "Requirement", "Evidence", "Gap", "Action"), (20, 27, 20, 13, 20)),
        ("Remediation tracker", ("Action", "Owner", "Due", "Priority", "Status"), (35, 18, 15, 15, 17)),
    ],
    "8.1": [
        ("Service objectives", ("Service / Journey", "SLI", "SLO", "Window", "Owner"), (25, 20, 18, 18, 19)),
        ("Error budget policy", ("Condition", "Budget state", "Required action", "Decision owner"), (23, 20, 38, 19)),
    ],
    "8.2": [
        ("Runbook procedures", ("Trigger", "Diagnostic", "Action / Command", "Verify", "Escalate"), (19, 24, 29, 16, 12)),
        ("Operational contacts", ("Service", "Primary", "Backup", "Escalation", "Hours"), (24, 19, 19, 23, 15)),
    ],
    "8.3": [
        ("Incident timeline", ("Time", "Event / Observation", "Actor", "Evidence"), (16, 42, 18, 24)),
        ("Corrective actions", ("Action", "Type", "Owner", "Due", "Verification"), (35, 14, 18, 13, 20)),
    ],
    "8.4": [
        ("Recovery objectives", ("Service", "Tier", "RTO", "RPO", "Recovery owner"), (25, 14, 15, 15, 31)),
        ("DR test plan", ("Scenario", "Scope", "Success criteria", "Date", "Owner"), (27, 22, 27, 12, 12)),
    ],
    "10.1": [
        ("Deployment sequence", ("Step", "Component", "Command / Action", "Verify", "Owner"), (10, 22, 34, 21, 13)),
        ("Rollback plan", ("Trigger", "Rollback action", "Data impact", "Verification", "Decision owner"), (18, 29, 18, 21, 14)),
    ],
    "10.2": [
        ("Go / No-Go checklist", ("Gate item", "Evidence", "Confirming role", "Result", "Comment"), (31, 25, 18, 12, 14)),
        ("Release decision", ("Decision", "Condition", "Decision owner", "Time"), (20, 41, 24, 15)),
    ],
    "10.3": [
        ("Pipeline stages", ("Stage", "Trigger", "Quality gate", "Artifact", "Owner"), (17, 20, 27, 20, 16)),
        ("Pipeline controls", ("Control", "Rule", "Failure action", "Evidence"), (24, 35, 25, 16)),
    ],
    "10.4": [
        ("Environment comparison", ("Dimension", "DEV", "SIT", "UAT", "PROD"), (24, 19, 19, 19, 19)),
        ("Configuration inventory", ("Variable / Resource", "Reference", "Owner", "Drift check", "Status"), (30, 24, 16, 17, 13)),
    ],
    "11.1": [
        ("Decision rights", ("Decision", "PM", "PMO", "SteerCo", "Threshold"), (32, 13, 13, 17, 25)),
        ("PMO service catalogue", ("Service", "Consumer", "Cadence", "Outcome", "Owner"), (25, 20, 17, 23, 15)),
    ],
    "11.2": [
        ("Gate scorecard", ("Criterion", "Weight", "Score", "Evidence", "Comment"), (30, 13, 13, 26, 18)),
        ("Gate decision record", ("Decision", "Condition / Action", "Owner", "Due", "Approval"), (18, 38, 17, 13, 14)),
    ],
    "11.3": [
        ("Portfolio dashboard", ("Project", "RAG", "% Complete", "Next milestone", "Leadership ask"), (20, 12, 16, 24, 28)),
        ("Cross-project conflicts", ("Conflict / Dependency", "Projects", "Impact", "Decision", "Owner"), (30, 20, 18, 20, 12)),
    ],
    "11.4": [
        ("Benefit register", ("Benefit", "Measure / Formula", "Baseline", "Target", "Business owner"), (24, 29, 14, 14, 19)),
        ("Benefit review plan", ("Review point", "Data source", "Decision if missed", "Owner"), (20, 25, 37, 18)),
    ],
}


DIAGRAMS = {
    "5.2": ("Sitemap / User Flow", ("[HOME]", "[SECTION]", "[SCREEN]", "[ACTION]"), "flowchart LR\n  HOME[HOME] --> SECTION[SECTION]\n  SECTION --> SCREEN[SCREEN]\n  SCREEN --> ACTION[ACTION]"),
    "7.1": ("Data Flow Diagram", ("[EXTERNAL]", "[BOUNDARY]", "[SYSTEM]", "[DATA STORE]"), "flowchart LR\n  EXT[External entity] --> BOUNDARY[Trust boundary]\n  BOUNDARY --> SYSTEM[System]\n  SYSTEM --> STORE[(Data store)]"),
    "10.3": ("CI/CD Pipeline", ("COMMIT", "BUILD", "TEST / SCAN", "DEPLOY", "VERIFY"), "flowchart LR\n  COMMIT --> BUILD --> TEST_SCAN[Test and Scan]\n  TEST_SCAN --> DEPLOY --> VERIFY"),
    "11.4": ("Benefit Map", ("OUTPUT", "OUTCOME", "BENEFIT", "STRATEGIC GOAL"), "flowchart LR\n  OUTPUT --> OUTCOME --> BENEFIT --> GOAL[Strategic goal]"),
}


def clean_markdown(value: str) -> str:
    value = re.sub(r"\*\*|`", "", value)
    return re.sub(r"\s+", " ", value).strip(" -")


def parse_prompt_pack(path: Path) -> dict[str, PromptRecord]:
    text = path.read_text(encoding="utf-8-sig")
    headings = list(re.finditer(r"^###\s+(\d+\.\d+)\.?\s+(.+)$", text, re.MULTILINE))
    records: dict[str, PromptRecord] = {}
    for index, match in enumerate(headings):
        code = match.group(1)
        raw_heading = clean_markdown(match.group(2))
        inline_pic = re.search(r"\s+[—–-]\s+PIC:\s*(.+)$", raw_heading, re.IGNORECASE)
        heading = raw_heading[:inline_pic.start()].strip() if inline_pic else raw_heading
        end = headings[index + 1].start() if index + 1 < len(headings) else len(text)
        block = text[match.end():end]
        pic_match = re.search(r"^\*\*PIC:\s*(.+?)\*\*", block, re.MULTILINE)
        pic = clean_markdown(pic_match.group(1)) if pic_match else (
            clean_markdown(inline_pic.group(1)) if inline_pic else "[ĐIỀN: PIC]"
        )
        fenced = re.search(r"```\s*\n([\s\S]*?)\n```", block)
        body = fenced.group(1) if fenced else block

        def bullets_between(start_pattern: str, end_pattern: str) -> tuple[str, ...]:
            start = re.search(start_pattern, body, re.IGNORECASE | re.MULTILINE)
            if not start:
                return ()
            tail = body[start.end():]
            end_match = re.search(end_pattern, tail, re.IGNORECASE | re.MULTILINE)
            area = tail[:end_match.start()] if end_match else tail
            items = []
            current = ""
            for line in area.splitlines():
                if re.match(r"^\s*(?:[-*]|\d+\.)\s+", line):
                    if current:
                        items.append(clean_markdown(current))
                    current = re.sub(r"^\s*(?:[-*]|\d+\.)\s+", "", line)
                elif current and line.strip():
                    current += " " + line.strip()
            if current:
                items.append(clean_markdown(current))
            return tuple(item for item in items if item)

        def parse_inputs() -> tuple[str, ...]:
            start = re.search(
                r"^(?:Input riêng|Thông tin đầu vào|Dữ liệu đầu vào(?:\s*\([^\n]+\))?|"
                r"User Story & Acceptance Criteria đầu vào|Thông tin lỗi quan sát được|"
                r"Dữ liệu kết quả test|Thông tin sự cố):\s*",
                body,
                re.IGNORECASE | re.MULTILINE,
            )
            if not start:
                return ()
            tail = body[start.end():]
            end_match = re.search(
                r"^(?:Cấu trúc bắt buộc|Cấu trúc|Yêu cầu|Phần 1|Nguyên tắc best practice)\s*[:–—]",
                tail,
                re.IGNORECASE | re.MULTILINE,
            )
            area = tail[:end_match.start()] if end_match else tail
            items = []
            for line in area.splitlines():
                line = line.strip()
                if not line:
                    continue
                line = re.sub(r"^[-*]\s+", "", line)
                value = clean_markdown(line)
                if value:
                    items.append(value)
            return tuple(items)

        inputs = parse_inputs()
        requirements_area = re.split(
            r"^Nguyên tắc best practice:\s*",
            body,
            maxsplit=1,
            flags=re.IGNORECASE | re.MULTILINE,
        )[0]
        section_items: list[str] = []
        current_section = ""
        group_prefix = ""
        for line in requirements_area.splitlines():
            stripped = line.strip()
            if re.match(r"^Phần\s+\d+\s*[–—-]", stripped, re.IGNORECASE):
                if current_section:
                    section_items.append(clean_markdown(current_section))
                    current_section = ""
                group_prefix = clean_markdown(stripped.rstrip(":"))
                continue
            numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
            if numbered:
                if current_section:
                    section_items.append(clean_markdown(current_section))
                value = numbered.group(1)
                current_section = f"{group_prefix}: {value}" if group_prefix else value
            elif current_section and stripped:
                current_section += " " + stripped
        if current_section:
            section_items.append(clean_markdown(current_section))
        sections = tuple(section_items)
        principles = bullets_between(r"^Nguyên tắc best practice:\s*", r"^Output:")
        output_match = re.search(r"^Output:\s*(.+)$", body, re.MULTILINE)
        output = clean_markdown(output_match.group(1)) if output_match else "File Word (.docx)"
        records[code] = PromptRecord(code, heading, pic, inputs, sections, principles, output)
    return records


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK, size: float = 9.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_run(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run(run, *, size: float = 10.5, color: str = INK, bold: bool = False, italic: bool = False, font: str = "Arial") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, weights: tuple[int, ...]) -> None:
    total = sum(weights)
    widths = [round(TABLE_WIDTH_DXA * weight / total) for weight in weights]
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(doc, title: str, columns: tuple[str, ...], weights: tuple[int, ...], rows: int = 3) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(title), size=11, color=POWER_DARK_RED, bold=True)
    table = doc.add_table(rows=1 + rows, cols=len(columns))
    table.style = "Table Grid"
    set_table_geometry(table, weights)
    set_repeat_header(table.rows[0])
    for cell, label in zip(table.rows[0].cells, columns):
        set_cell_shading(cell, POWER_RED)
        set_cell_text(cell, label, bold=True, color=WHITE, size=8.5)
    for row_index, row in enumerate(table.rows[1:], 1):
        for col_index, cell in enumerate(row.cells):
            if row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            placeholder = "[ĐIỀN]"
            if col_index == 0:
                placeholder = f"[ĐIỀN {row_index}]"
            set_cell_text(cell, placeholder, color=MUTED, size=8.5)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_heading(doc, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True


def add_body(doc, text: str, *, bold_lead: str | None = None, italic: bool = False) -> None:
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_run(paragraph.add_run(bold_lead), bold=True)
        set_run(paragraph.add_run(text[len(bold_lead):]), italic=italic)
    else:
        set_run(paragraph.add_run(text), italic=italic)


def add_bullet(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run(paragraph.add_run(text))


def add_callout(doc, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(9)
    paragraph.paragraph_format.right_indent = Pt(9)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_RED)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), POWER_RED)
    borders.append(left)
    p_pr.append(borders)
    set_run(paragraph.add_run(f"{label}: "), color=POWER_DARK_RED, bold=True)
    set_run(paragraph.add_run(text), color=INK)


def add_field(paragraph, instruction: str, fallback: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    run.extend((begin, instr, separate, text, end))


def configure_styles(doc) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, before, after, color in (
        (1, 16, 18, 10, POWER_DARK_RED),
        (2, 13, 14, 7, POWER_RED),
        (3, 11.5, 10, 5, INK),
    ):
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_page(doc, role_label: str, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(19.05)
    section.bottom_margin = Mm(19.05)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(2)
    set_run(hp.add_run("POWER HOME  "), size=8.5, color=POWER_RED, bold=True)
    set_run(hp.add_run(f"|  {role_label}"), size=8.5, color=MUTED, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.27), WD_TAB_ALIGNMENT.RIGHT)
    set_run(p.add_run("MẬT (CONFIDENTIAL) · LƯU HÀNH NỘI BỘ"), size=8, color=MUTED)
    set_run(p.add_run("\tTrang "), size=8, color=MUTED)
    add_field(p, "PAGE", "1")

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_cover(doc, record: PromptRecord, role_label: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run(f"POWER HOME · {role_label}"), size=10, color=POWER_RED, bold=True)

    title = record.heading
    title = re.sub(r"^Prompt sinh\s+", "", title, flags=re.IGNORECASE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(title.upper()), size=25, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(26)
    set_run(p.add_run("MASTER TEMPLATE · v1.0 · DRAFT"), size=11, color=POWER_DARK_RED, bold=True)

    meta = doc.add_table(rows=6, cols=2)
    meta.style = "Table Grid"
    set_table_geometry(meta, (27, 73))
    set_repeat_header(meta.rows[0])
    set_cell_shading(meta.rows[0].cells[0], POWER_RED)
    set_cell_shading(meta.rows[0].cells[1], POWER_RED)
    set_cell_text(meta.rows[0].cells[0], "Trường", bold=True, color=WHITE)
    set_cell_text(meta.rows[0].cells[1], "Giá trị kiểm soát", bold=True, color=WHITE)
    rows = (
        ("Mã tài liệu", f"[ĐIỀN: <TÊN DỰ ÁN>-{record.code.replace('.', '')}-v1.0-YYYYMMDD]"),
        ("PIC duy nhất", record.pic),
        ("Trạng thái", "Draft → In Review → Approved → Superseded"),
        ("Ngày hiệu lực", "[ĐIỀN: dd/mm/yyyy]"),
        ("Phân loại", "MẬT (CONFIDENTIAL) · Lưu hành nội bộ"),
    )
    for row, (label, value) in zip(meta.rows[1:], rows):
        set_cell_shading(row.cells[0], LIGHT_RED)
        set_cell_text(row.cells[0], label, bold=True, color=POWER_DARK_RED)
        set_cell_text(row.cells[1], value)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    add_callout(
        doc,
        "Nguyên tắc sử dụng",
        "Đây là template dùng chung, không gắn với dự án cụ thể. Giữ nguyên mọi ô [ĐIỀN...] cho đến khi PIC có dữ liệu được xác nhận; không tự suy diễn bối cảnh hay số liệu.",
    )
    for item in (
        "Chỉ bản Approved được dùng làm căn cứ thực thi hoặc ra quyết định.",
        "Các role khác đóng góp qua review/input; chỉ PIC cập nhật bản chính.",
        "Mọi số liệu chưa xác nhận phải ghi [CẦN XÁC NHẬN].",
    ):
        add_bullet(doc, item)
    doc.add_page_break()


def add_front_matter(doc, record: PromptRecord) -> None:
    add_heading(doc, "Kiểm soát tài liệu", 1)
    add_table(
        doc,
        "Revision History",
        ("Version", "Ngày", "Người sửa (PIC)", "Nội dung thay đổi", "Người duyệt"),
        (12, 14, 23, 33, 18),
        rows=2,
    )
    table = doc.tables[-1]
    values = ("1.0", "[ĐIỀN]", record.pic, "Bản template đầu tiên", "[ĐIỀN]")
    for cell, value in zip(table.rows[1].cells, values):
        set_cell_text(cell, value, size=8.5)

    add_heading(doc, "Mục lục", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "Mở tài liệu trong Word và chọn Update Field để cập nhật mục lục.")
    add_callout(doc, "Output yêu cầu", record.output)
    doc.add_page_break()


def section_title(requirement: str, ordinal: int) -> str:
    text = re.sub(r"^\d+\.\s*", "", requirement)
    candidate = re.split(r"\s+[—–-]\s+|:\s+", text, maxsplit=1)[0]
    candidate = candidate.strip().rstrip(".")
    if len(candidate) > 80:
        candidate = " ".join(candidate.split()[:10])
    return f"{ordinal}. {candidate}"


def draw_flow_image(temp_dir: Path, code: str, title: str, nodes: tuple[str, ...]) -> Path:
    width, height = 1500, 360
    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 40)
        node_font = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 26)
    except OSError:
        title_font = ImageFont.load_default()
        node_font = ImageFont.load_default()
    draw.text((60, 35), title, fill="#25324A", font=title_font)
    gap = 35
    usable = width - 120 - gap * (len(nodes) - 1)
    node_width = usable // len(nodes)
    top, bottom = 145, 255
    centers = []
    for index, node in enumerate(nodes):
        left = 60 + index * (node_width + gap)
        right = left + node_width
        draw.rounded_rectangle((left, top, right, bottom), radius=20, fill="#FDEDEE", outline="#DA2128", width=4)
        box = draw.textbbox((0, 0), node, font=node_font)
        x = left + (node_width - (box[2] - box[0])) / 2
        y = top + (bottom - top - (box[3] - box[1])) / 2 - 3
        draw.text((x, y), node, fill="#A71920", font=node_font)
        centers.append((left, right))
    for index in range(len(centers) - 1):
        start = (centers[index][1] + 4, (top + bottom) // 2)
        end = (centers[index + 1][0] - 4, (top + bottom) // 2)
        draw.line((start, end), fill="#667085", width=5)
        draw.polygon(((end[0], end[1]), (end[0] - 16, end[1] - 9), (end[0] - 16, end[1] + 9)), fill="#667085")
    output = temp_dir / f"diagram-{code.replace('.', '-')}.png"
    image.save(output)
    return output


def add_diagram(doc, temp_dir: Path, code: str) -> None:
    title, nodes, mermaid = DIAGRAMS[code]
    add_heading(doc, title, 2)
    image_path = draw_flow_image(temp_dir, code, title, nodes)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    inline = run.add_picture(str(image_path), width=Inches(6.15))
    inline._inline.docPr.set("descr", f"{title} placeholder flow; replace nodes with approved project content")
    add_body(doc, "Hình minh hoạ là khung placeholder. PIC phải thay nhãn bằng nội dung dự án đã xác nhận.", italic=True)
    add_heading(doc, f"Phụ lục · Mã Mermaid cho {title}", 2)
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Pt(9)
    code.paragraph_format.right_indent = Pt(9)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_GRAY)
    code._p.get_or_add_pPr().append(shading)
    run = code.add_run(mermaid)
    set_run(run, size=8.5, color=INK, font="Consolas")


def build_document(record: PromptRecord, role: str, role_label: str, destination: Path, curator) -> None:
    doc = Document()
    configure_styles(doc)
    title = re.sub(r"^Prompt sinh\s+", "", record.heading, flags=re.IGNORECASE)
    configure_page(doc, role_label, title)
    doc.core_properties.title = title
    doc.core_properties.subject = f"Power Home {role_label} master template"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = "Generated from the reviewed role prompt pack; PIC review required."
    add_cover(doc, record, role_label)
    add_front_matter(doc, record)

    add_heading(doc, "1. Thông tin đầu vào", 1)
    inputs = record.inputs or ("[ĐIỀN: dữ liệu đầu vào đã được PIC xác nhận]",)
    table = doc.add_table(rows=1 + len(inputs), cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, (34, 48, 18))
    set_repeat_header(table.rows[0])
    for cell, label in zip(table.rows[0].cells, ("Đầu vào bắt buộc", "Giá trị / Tham chiếu", "Trạng thái")):
        set_cell_shading(cell, POWER_RED)
        set_cell_text(cell, label, bold=True, color=WHITE, size=8.5)
    for row, item in zip(table.rows[1:], inputs):
        set_cell_text(row.cells[0], item, size=8.5)
        set_cell_text(row.cells[1], "[ĐIỀN: nội dung hoặc đường dẫn nguồn]", color=MUTED, size=8.5)
        set_cell_text(row.cells[2], "[CẦN XÁC NHẬN]", color=POWER_DARK_RED, size=8.5)

    add_heading(doc, "2. Nội dung tài liệu", 1)
    for ordinal, requirement in enumerate(record.sections, 1):
        add_heading(doc, section_title(requirement, ordinal), 2)
        add_callout(doc, "Yêu cầu từ prompt", requirement)
        add_body(doc, "[ĐIỀN: nội dung đã được PIC xác nhận, kèm nguồn hoặc minh chứng khi áp dụng]")
        add_body(doc, "Quyết định / open point: [ĐIỀN hoặc ghi Không áp dụng kèm lý do]")

    add_heading(doc, "3. Biểu mẫu làm việc", 1)
    for sheet_title, columns, weights in WORKSHEETS[record.code]:
        add_table(doc, sheet_title, columns, weights, rows=3)

    if record.code in DIAGRAMS:
        with tempfile.TemporaryDirectory(prefix="handbook-diagram-") as temp:
            add_diagram(doc, Path(temp), record.code)

    add_heading(doc, "4. Nguyên tắc chất lượng", 1)
    for principle in record.principles:
        add_bullet(doc, principle)
    add_bullet(doc, "PIC phải review, chịu trách nhiệm nội dung và chuyển trạng thái theo vòng đời tài liệu.")
    add_bullet(doc, "Không để lại secret thật; chỉ ghi tên biến và tham chiếu secret manager khi cần.")

    add_heading(doc, "5. Review & phê duyệt", 1)
    add_table(
        doc,
        "Sign-off",
        ("Vai trò", "Họ tên", "Kết luận", "Ngày", "Điều kiện / Ghi chú"),
        (20, 20, 18, 14, 28),
        rows=3,
    )

    destination.parent.mkdir(parents=True, exist_ok=True)
    PREVIEW_ROOT.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="handbook-role-template-") as temp:
        intermediate = Path(temp) / destination.name
        doc.save(intermediate)
        curator.sanitized_package(intermediate, destination)

    final_doc = Document(destination)
    preview_name = destination.stem + ".html"
    preview_html = curator.render_preview_page(final_doc, f"{role}/{destination.name}")
    (PREVIEW_ROOT / preview_name).write_text(preview_html, encoding="utf-8", newline="\n")
    print(f"[OK] {record.code} {destination.relative_to(ROOT)}")
    print(f"[OK] preview {(PREVIEW_ROOT / preview_name).relative_to(ROOT)}")


def load_curator_module():
    path = ROOT / "scripts" / "curate-documents.py"
    spec = importlib.util.spec_from_file_location("handbook_curator", path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot import {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--prompt-file", required=True, type=Path)
    parser.add_argument("--role", required=True, choices=tuple(ROLE_SPECS))
    parser.add_argument("--force", action="store_true", help="Replace files previously generated by this script")
    args = parser.parse_args()
    records = parse_prompt_pack(args.prompt_file)
    role_spec = ROLE_SPECS[args.role]
    curator = load_curator_module()
    for code in role_spec["codes"]:
        if code not in records:
            raise ValueError(f"Prompt {code} was not found in {args.prompt_file}")
        destination = OUTPUT_ROOT / args.role / role_spec["files"][code]
        if destination.exists() and not args.force:
            raise FileExistsError(f"Refusing to overwrite existing template: {destination}")
        build_document(records[code], args.role, role_spec["label"], destination, curator)
    print(f"[DONE] role={args.role} templates={len(role_spec['codes'])}")


if __name__ == "__main__":
    main()
