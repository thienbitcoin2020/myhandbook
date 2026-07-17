#!/usr/bin/env python3
"""Build the reviewed, publishable DOCX template library.

The source files under ``Template/`` are working material and are intentionally
not published. This script creates role-owned, stable-name copies under
``assets/templates/`` after removing duplicated document-control boilerplate,
sample data, and role rules that do not belong in the artifact itself.

Requires: python-docx and lxml.
"""

from __future__ import annotations

import argparse
import re
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import Callable, Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
SOURCE_ROOT = ROOT / "Template"
DEFAULT_OUTPUT_ROOT = ROOT / "assets" / "templates"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"

REMOVED_PACKAGE_PARTS = {
    "docProps/custom.xml",
    "word/comments.xml",
    "word/commentsExtended.xml",
    "word/commentsExtensible.xml",
    "word/commentsIds.xml",
    "word/people.xml",
}

FORBIDDEN_PART_PATTERNS = (
    "vbaproject.bin",
    "/activex/",
    "/embeddings/",
    "oleobject",
    "altchunk",
    "afchunk",
)


def body_children(document: DocumentObject) -> list:
    return list(document.element.body.iterchildren())


def body_text(element) -> str:
    return " ".join("".join(element.itertext()).split())


def require_body_text(document: DocumentObject, index: int, expected: str) -> None:
    children = body_children(document)
    if index >= len(children) or expected not in body_text(children[index]):
        actual = body_text(children[index])[:120] if index < len(children) else "<missing>"
        raise ValueError(
            f"Source structure changed at body index {index}: expected {expected!r}, got {actual!r}"
        )


def remove_body_indices(document: DocumentObject, indices: Iterable[int]) -> None:
    children = body_children(document)
    for index in sorted(set(indices), reverse=True):
        if index < 0 or index >= len(children):
            raise IndexError(f"Body index {index} is outside 0..{len(children) - 1}")
        document.element.body.remove(children[index])


def replace_body_paragraph(document: DocumentObject, index: int, text: str) -> None:
    element = body_children(document)[index]
    if element.tag != qn("w:p"):
        raise ValueError(f"Body index {index} is not a paragraph")
    paragraph = Paragraph(element, document)
    paragraph_property = element.pPr
    for child in list(element):
        if child is not paragraph_property:
            element.remove(child)
    paragraph.add_run(text)


def table_header(table: Table) -> str:
    if not table.rows:
        return ""
    return " | ".join(cell.text.strip() for cell in table.rows[0].cells)


def find_table(document: DocumentObject, first_cell: str) -> Table:
    for table in document.tables:
        if table.rows and table.rows[0].cells[0].text.strip() == first_cell:
            return table
    raise ValueError(f"Could not find table whose first cell is {first_cell!r}")


def keep_labeled_rows(table: Table, labels: Iterable[str], keep_header: bool = False) -> None:
    allowed = set(labels)
    for index in range(len(table.rows) - 1, -1, -1):
        label = table.rows[index].cells[0].text.strip()
        if (keep_header and index == 0) or label in allowed:
            continue
        table._tbl.remove(table.rows[index]._tr)


def replace_labeled_value(table: Table, label: str, value: str) -> None:
    for row in table.rows:
        if row.cells[0].text.strip() == label:
            row.cells[1].text = value
            return
    raise ValueError(f"Could not find row {label!r} in table {table_header(table)!r}")


def reset_data_rows(table: Table, rows: list[list[str]]) -> None:
    expected_columns = len(table.columns)
    for values in rows:
        if len(values) != expected_columns:
            raise ValueError(
                f"Row for table {table_header(table)!r} has {len(values)} cells; expected {expected_columns}"
            )
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = value


def trim_epic(document: DocumentObject) -> None:
    require_body_text(document, 5, "Chủ trì (PIC)")
    require_body_text(document, 7, "Cách dùng & phân vai")
    require_body_text(document, 44, "US-02")
    require_body_text(document, 54, "Lịch sử thay đổi")

    replacements = {
        13: "1 · Input epic cần phân rã",
        15: "2 · Nguyên tắc phân rã & viết AC",
        16: "2.1 · Chuẩn INVEST — mỗi story phải qua đủ 6 cửa",
        19: "2.2 · Story quá lớn → tách theo các mẫu sau",
        26: "2.3 · Thin vertical slice — cắt dọc, giao sớm",
        30: "2.4 · Từ cần tránh trong Acceptance Criteria",
        33: "3 · Backlog tổng quan — thứ tự giao giá trị",
        34: (
            "Sắp story theo thứ tự giao giá trị sớm nhất, không theo thứ tự kỹ thuật. "
            "Chỉ tạo các story cần thiết để epic có thể được lập kế hoạch và kiểm chứng."
        ),
        39: "Size S/M/L là tương đối giữa các story trong epic; story L → quay lại mục 2.2 để tách tiếp.",
        40: "4 · Chi tiết User Story — mỗi story một block",
        41: (
            "Mẫu cung cấp một block US-01. Nhân bản block khi cần và chỉ giữ các story thực sự thuộc epic; "
            "mỗi story phải có giá trị, tiêu chí chấp nhận và quan hệ phụ thuộc rõ ràng."
        ),
    }
    for index, text in replacements.items():
        replace_body_paragraph(document, index, text)

    remove_body_indices(
        document,
        list(range(5, 13)) + list(range(44, 54)) + [54, 55],
    )


def trim_product_vision(document: DocumentObject) -> None:
    require_body_text(document, 5, "Chủ trì (PIC)")
    require_body_text(document, 7, "Cách dùng template")
    require_body_text(document, 67, "Lịch sử thay đổi")

    metadata = document.tables[0]
    keep_labeled_rows(metadata, {"Tên sản phẩm", "Chu kỳ review roadmap"})
    replace_body_paragraph(document, 7, "1 · Input đầu vào")
    remove_body_indices(document, [5, *range(8, 13), 67, 68])


def trim_rtm(document: DocumentObject) -> None:
    require_body_text(document, 8, "Kiểm soát tài liệu")
    require_body_text(document, 49, "Quy trình cập nhật RTM")

    metadata = document.tables[0]
    keep_labeled_rows(
        metadata,
        {"Mã tài liệu", "Phạm vi", "Trạng thái", "Ngày cập nhật gần nhất", "Mức bảo mật"},
    )
    replace_labeled_value(metadata, "Mã tài liệu", "[ĐIỀN: RTM-<DỰ ÁN>-<PHẠM VI>]")

    matrix = find_table(document, "BR ID")
    for index in range(len(matrix.rows) - 1, 0, -1):
        if "[TRỐNG — gap]" in " ".join(cell.text for cell in matrix.rows[index].cells):
            matrix._tbl.remove(matrix.rows[index]._tr)

    remove_body_indices(document, [*range(8, 16), 33, 34, *range(49, 57)])


def trim_srs(document: DocumentObject) -> None:
    require_body_text(document, 8, "Kiểm soát tài liệu")
    require_body_text(document, 65, "Feature 2")
    require_body_text(document, 111, "Tóm tắt BR chính")

    metadata = document.tables[0]
    keep_labeled_rows(metadata, {"Mã tài liệu", "Module / Phân hệ", "Trạng thái", "Mức bảo mật"})
    replace_labeled_value(metadata, "Mã tài liệu", "[ĐIỀN: SRS-<DỰ ÁN>-<MODULE>]")
    replace_body_paragraph(
        document,
        88,
        (
            "BA và PO làm rõ nhu cầu chất lượng; Solution Architect thiết kế giải pháp đáp ứng; "
            "Security, Ops và QC cùng xác định cách kiểm chứng. Ghi yêu cầu đo được tại đây hoặc "
            "tham chiếu nguồn yêu cầu phi chức năng chuẩn của dự án."
        ),
    )
    replace_body_paragraph(
        document,
        73,
        (
            "Nhân bản cấu trúc 3.1 cho từng feature và đánh mã FR liên tục trong toàn tài liệu "
            "để giữ traceability duy nhất."
        ),
    )
    remove_body_indices(document, [*range(8, 18), *range(65, 73), *range(111, 115)])


def trim_use_case(document: DocumentObject) -> None:
    require_body_text(document, 8, "Kiểm soát tài liệu")
    require_body_text(document, 53, "UC-002")
    require_body_text(document, 78, "Nhân bản toàn bộ section UC")

    metadata = document.tables[0]
    keep_labeled_rows(metadata, {"Mã tài liệu", "Chức năng / Module", "Trạng thái", "Mức bảo mật"})
    replace_labeled_value(metadata, "Mã tài liệu", "[ĐIỀN: UCS-<DỰ ÁN>-<MODULE>]")
    replace_body_paragraph(
        document,
        78,
        (
            "Nhân bản mục UC-001 cho mỗi use case mới; giữ mã UC duy nhất và cập nhật bảng "
            "Use Case Overview ở mục 1.4."
        ),
    )
    remove_body_indices(document, [*range(8, 16), *range(53, 78)])


def trim_project_management_plan(document: DocumentObject) -> None:
    require_body_text(document, 1, "Template dùng chung")
    require_body_text(document, 4, "Thông tin tài liệu")
    require_body_text(document, 7, "Revision History")

    metadata = document.tables[1]
    keep_labeled_rows(
        metadata,
        {
            "Trường",
            "Tên dự án",
            "Mã dự án",
            "Project Charter tham chiếu",
            "Phương pháp thực thi",
            "Trạng thái tài liệu",
            "Chuẩn tham chiếu",
        },
        keep_header=True,
    )
    replace_labeled_value(
        metadata,
        "Chuẩn tham chiếu",
        "PMBOK® Guide — Seventh Edition (tài liệu tham chiếu của mẫu)",
    )
    replace_body_paragraph(document, 4, "Thông tin kế hoạch")
    remove_body_indices(document, [1, 2, 3, 7, 8, 9])


def trim_raid(document: DocumentObject) -> None:
    require_body_text(document, 1, "Template dùng chung")
    require_body_text(document, 7, "Update Log")
    require_body_text(document, 10, "Quy trình đóng góp")

    metadata = document.tables[1]
    keep_labeled_rows(
        metadata,
        {
            "Trường",
            "Tên dự án",
            "Mã dự án",
            "Khẩu vị rủi ro của Sponsor",
            "Tần suất review định kỳ",
            "Vị trí lưu bản chính (single source of truth)",
            "Tài liệu tham chiếu",
            "Chuẩn tham chiếu",
        },
        keep_header=True,
    )
    replace_labeled_value(
        metadata,
        "Chuẩn tham chiếu",
        "PMBOK® Guide — Seventh Edition (tài liệu tham chiếu của mẫu)",
    )

    risk_table = find_table(document, "ID")
    if "Mô tả (nguyên nhân" not in table_header(risk_table):
        raise ValueError("Unexpected first ID table; RAID risk table could not be identified")
    reset_data_rows(
        risk_table,
        [
            [
                f"R-{index:03d}",
                "[ĐIỀN: nguyên nhân → sự kiện → tác động]",
                "[ĐIỀN]",
                "[1–5]",
                "[1–5]",
                "[P×I]",
                "[Avoid / Mitigate / Transfer / Accept]",
                "[ĐIỀN: mitigation / contingency]",
                "[ĐIỀN tên]",
                "[ĐIỀN điều kiện kích hoạt]",
                "Open",
            ]
            for index in range(1, 5)
        ],
    )

    assumption_table = find_table(document, "ID")
    id_tables = [table for table in document.tables if table.rows and table.rows[0].cells[0].text.strip() == "ID"]
    assumption_table = next(table for table in id_tables if "Mô tả giả định" in table_header(table))
    issue_table = next(table for table in id_tables if "Mô tả vấn đề" in table_header(table))
    dependency_table = next(table for table in id_tables if "Mô tả phụ thuộc" in table_header(table))

    reset_data_rows(
        assumption_table,
        [
            [f"A-{index:03d}", "[ĐIỀN giả định]", "[ĐIỀN tác động]", "[ĐIỀN hạn validate]", "[ĐIỀN tên]"]
            for index in range(1, 4)
        ],
    )
    reset_data_rows(
        issue_table,
        [
            [
                f"I-{index:03d}",
                "[ĐIỀN vấn đề hiện hữu]",
                "[Critical / High / Medium / Low]",
                "[ĐIỀN hành động xử lý]",
                "[ĐIỀN tên]",
                "[ĐIỀN]",
                "Open",
            ]
            for index in range(1, 4)
        ],
    )
    reset_data_rows(
        dependency_table,
        [
            [
                f"D-{index:03d}",
                "[ĐIỀN phụ thuộc]",
                "[Internal / External]",
                "[ĐIỀN bên liên quan]",
                "[ĐIỀN due date]",
                "Open",
            ]
            for index in range(1, 4)
        ],
    )

    replace_body_paragraph(document, 4, "Ngữ cảnh RAID Log")
    replace_body_paragraph(
        document,
        24,
        (
            "Mô tả theo cấu trúc « Do <nguyên nhân>, có thể xảy ra <sự kiện>, dẫn đến <tác động> ». "
            "Chấm P/I trong risk workshop và cập nhật khi có dữ liệu mới."
        ),
    )
    replace_body_paragraph(
        document,
        28,
        (
            "Chỉ ghi các giả định có ảnh hưởng đến kế hoạch hoặc quyết định; mỗi giả định phải có "
            "owner và hạn xác nhận."
        ),
    )
    remove_body_indices(document, [1, 2, 3, *range(7, 13)])


def trim_status_report(document: DocumentObject) -> None:
    require_body_text(document, 1, "Template 1 trang")
    metadata = document.tables[0]
    while len(metadata.rows) > 1:
        metadata._tbl.remove(metadata.rows[-1]._tr)
    metadata.rows[0].cells[2].text = "Ngày chốt số liệu: [ĐIỀN]"
    remove_body_indices(document, [1])


def trim_project_charter(document: DocumentObject) -> None:
    require_body_text(document, 4, "Ghi chú sử dụng template")
    require_body_text(document, 81, "Nhật ký thay đổi")

    metadata = document.tables[1]
    keep_labeled_rows(metadata, {"Tên Dự Án", "Tổ Chức", "Giai đoạn", "Thời gian triển khai"})
    remove_body_indices(document, [4, 5, 81, 82])


def trim_adr(document: DocumentObject) -> None:
    require_body_text(document, 2, "Hướng dẫn sử dụng")
    require_body_text(document, 10, "Vòng đời trạng thái")
    require_body_text(document, 50, "Bản Markdown")
    remove_body_indices(document, [*range(1, 13), *range(50, 94)])


def trim_solution_architecture(document: DocumentObject) -> None:
    require_body_text(document, 6, "Kiểm soát tài liệu")
    require_body_text(document, 88, "Phụ lục A. Mã nguồn Mermaid")

    metadata = document.tables[0]
    keep_labeled_rows(
        metadata,
        {"Mã dự án", "Trạng thái", "Người review", "Người phê duyệt", "Phạm vi bảo mật"},
        keep_header=True,
    )
    remove_body_indices(document, [6, 7, 8, *range(88, 219)])


CURATIONS: tuple[tuple[str, str, Callable[[DocumentObject], None]], ...] = (
    ("BA/RTM_Template_BABOK.docx", "ba/requirements-traceability-matrix.docx", trim_rtm),
    ("BA/SRS_Template_ISO29148.docx", "ba/software-requirements-specification.docx", trim_srs),
    ("BA/UseCase_Specification_Template.docx", "ba/use-case-specification.docx", trim_use_case),
    (
        "PM/Project_Management_Plan_Template_PMBOK7.docx",
        "pm/project-management-plan.docx",
        trim_project_management_plan,
    ),
    ("PM/RAID_Log_Template_PMBOK7.docx", "pm/raid-log.docx", trim_raid),
    ("PM/Status_Report_Template_1page.docx", "pm/status-report-one-page.docx", trim_status_report),
    ("PM/Template_De_cuong_Du_an_Project_Charter.docx", "pm/project-charter.docx", trim_project_charter),
    ("PO/Epic_to_UserStories_Template.docx", "po/epic-to-user-stories.docx", trim_epic),
    ("PO/Product_Vision_Roadmap_Template.docx", "po/product-vision-roadmap.docx", trim_product_vision),
    ("SA/ADR-Template.docx", "sa/architecture-decision-record.docx", trim_adr),
    (
        "SA/Solution-Architecture-Document-Template.docx",
        "sa/solution-architecture-document.docx",
        trim_solution_architecture,
    ),
)


def scrub_story_xml(root: etree._Element) -> None:
    for element in list(root.iter()):
        for attribute in list(element.attrib):
            if attribute.startswith(f"{{{W_NS}}}rsid"):
                del element.attrib[attribute]
    for local_name in ("commentRangeStart", "commentRangeEnd", "commentReference"):
        for element in root.findall(f".//{{{W_NS}}}{local_name}"):
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)


def sanitized_package(source: Path, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(source, "r") as archive:
        names = archive.namelist()
        for name in names:
            normalized = "/" + name.replace("\\", "/").lower().lstrip("/")
            if any(pattern in normalized for pattern in FORBIDDEN_PART_PATTERNS):
                raise ValueError(f"Forbidden OOXML part in {source.name}: {name}")

        overrides: dict[str, bytes] = {}
        for name in names:
            if not name.endswith((".xml", ".rels")):
                continue
            root = etree.fromstring(archive.read(name))
            if name.endswith(".rels"):
                for relationship in list(root.findall(f"{{{REL_NS}}}Relationship")):
                    if (relationship.get("TargetMode") or "").lower() == "external":
                        raise ValueError(f"External relationship in {source.name}: {name}")
                    target = (relationship.get("Target") or "").replace("\\", "/")
                    if target.endswith("docProps/custom.xml") or "comments" in target or target.endswith("people.xml"):
                        root.remove(relationship)
                overrides[name] = etree.tostring(
                    root, xml_declaration=True, encoding="UTF-8", standalone="yes"
                )
                continue

            if name == "docProps/core.xml":
                for element in root.findall(f".//{{{DC_NS}}}creator"):
                    element.text = ""
                for element in root.findall(f".//{{{CP_NS}}}lastModifiedBy"):
                    element.text = ""
            if name == "[Content_Types].xml":
                for override in list(root.findall(f"{{{CT_NS}}}Override")):
                    part_name = (override.get("PartName") or "").lstrip("/")
                    if part_name in REMOVED_PACKAGE_PARTS or "comments" in part_name or part_name == "word/people.xml":
                        root.remove(override)
            if name.startswith("word/"):
                tracked = root.findall(f".//{{{W_NS}}}ins") + root.findall(f".//{{{W_NS}}}del")
                tracked += root.findall(f".//{{{W_NS}}}moveFrom") + root.findall(f".//{{{W_NS}}}moveTo")
                if tracked:
                    raise ValueError(f"Tracked changes must be resolved before publishing: {source.name} ({name})")
                scrub_story_xml(root)
            overrides[name] = etree.tostring(
                root, xml_declaration=True, encoding="UTF-8", standalone="yes"
            )

        with zipfile.ZipFile(destination, "w", zipfile.ZIP_DEFLATED) as output:
            for info in archive.infolist():
                name = info.filename
                if name in REMOVED_PACKAGE_PARTS or "word/comments" in name or name == "word/people.xml":
                    continue
                output.writestr(info, overrides.get(name, archive.read(name)))


def curate(source: Path, destination: Path, transform: Callable[[DocumentObject], None]) -> None:
    if not source.is_file():
        raise FileNotFoundError(source)
    document = Document(source)
    transform(document)
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    with tempfile.TemporaryDirectory(prefix="handbook-docx-") as temp_dir:
        intermediate = Path(temp_dir) / destination.name
        document.save(intermediate)
        sanitized_package(intermediate, destination)
    # Re-open the final package so generation fails immediately on structural corruption.
    Document(destination)


# ── Confluence MHTML sources ────────────────────────────────────────────
# BRD/FRD/PRD arrived as Confluence page exports saved with a .doc extension
# (MIME multipart + quoted-printable HTML — Word opens them, python-docx does
# not). They are parsed into structured blocks, trimmed by the same editorial
# rules as the native DOCX sources (no change history, no static TOC, no
# duplicated page title), then emitted as fresh DOCX packages through the
# same sanitizer.

import email
import email.policy
from html.parser import HTMLParser


class _ConfluenceBlockParser(HTMLParser):
    """Flatten Confluence export HTML into heading/para/list/table blocks."""

    _HEADINGS = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}
    _SKIP = {"script", "style", "head", "title"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[dict] = []
        self._skip_depth = 0
        self._heading: int | None = None
        self._runs: list[tuple[str, bool, bool]] = []
        self._bold = 0
        self._italic = 0
        self._in_list_item = False
        self._table: list[list[str]] | None = None
        self._row: list[str] | None = None
        self._cell: list[str] | None = None

    # ── text assembly ────────────────────────────────────────────────
    def _flush_paragraph(self, kind: str = "para") -> None:
        runs = [(t, b, i) for t, b, i in self._runs if t.strip()]
        if runs:
            block = {"kind": kind, "runs": runs}
            if kind == "heading":
                block["level"] = self._heading or 2
            self.blocks.append(block)
        self._runs = []

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP:
            self._skip_depth += 1
            return
        if self._skip_depth:
            return
        if tag in self._HEADINGS:
            self._flush_paragraph("li" if self._in_list_item else "para")
            self._heading = self._HEADINGS[tag]
        elif tag == "table" and self._table is None:
            self._flush_paragraph()
            self._table = []
        elif tag == "tr" and self._table is not None:
            self._row = []
        elif tag in ("td", "th") and self._row is not None:
            self._cell = []
        elif tag == "li":
            self._flush_paragraph()
            self._in_list_item = True
        elif tag in ("p", "div", "br") and self._table is None:
            self._flush_paragraph("li" if self._in_list_item else "para")
        elif tag in ("strong", "b"):
            self._bold += 1
        elif tag in ("em", "i"):
            self._italic += 1

    def handle_endtag(self, tag):
        if tag in self._SKIP:
            self._skip_depth = max(0, self._skip_depth - 1)
            return
        if self._skip_depth:
            return
        if tag in self._HEADINGS:
            self._flush_paragraph("heading")
            self._heading = None
        elif tag == "table" and self._table is not None:
            width = max((len(row) for row in self._table), default=0)
            rows = [row + [""] * (width - len(row)) for row in self._table if any(c.strip() for c in row)]
            if rows and width:
                self.blocks.append({"kind": "table", "rows": rows})
            self._table = None
        elif tag == "tr" and self._table is not None and self._row is not None:
            self._table.append(self._row)
            self._row = None
        elif tag in ("td", "th") and self._cell is not None:
            text = " ".join("".join(self._cell).split())
            if self._row is not None:
                self._row.append(text)
            self._cell = None
        elif tag == "li":
            self._flush_paragraph("li")
            self._in_list_item = False
        elif tag in ("p", "div"):
            self._flush_paragraph("li" if self._in_list_item else "para")
        elif tag in ("strong", "b"):
            self._bold = max(0, self._bold - 1)
        elif tag in ("em", "i"):
            self._italic = max(0, self._italic - 1)

    def handle_data(self, data):
        if self._skip_depth:
            return
        text = data.replace("\xa0", " ")
        if self._cell is not None:
            self._cell.append(text)
            return
        if text.strip() or (self._runs and not text.isspace()):
            collapsed = " ".join(text.split())
            if text[:1].isspace() and self._runs:
                collapsed = " " + collapsed
            if text[-1:].isspace():
                collapsed += " "
            if collapsed:
                self._runs.append((collapsed, self._bold > 0, self._italic > 0))


def _block_text(block: dict) -> str:
    if block["kind"] == "table":
        return " ".join(cell for row in block["rows"] for cell in row)
    return "".join(text for text, _, _ in block.get("runs", []))


def parse_confluence_export(source: Path) -> list[dict]:
    message = email.message_from_bytes(source.read_bytes(), policy=email.policy.default)
    html = None
    for part in message.walk():
        if part.get_content_type() == "text/html":
            html = part.get_content()
            break
    if html is None:
        raise ValueError(f"{source.name}: no text/html part in the Confluence export")
    parser = _ConfluenceBlockParser()
    parser.feed(html)
    parser._flush_paragraph()
    return parser.blocks


def drop_block_section(blocks: list[dict], title_marker: str) -> list[dict]:
    """Remove a heading and everything under it (until a same/higher heading)."""
    result: list[dict] = []
    index = 0
    found = False
    while index < len(blocks):
        block = blocks[index]
        if (not found and block["kind"] == "heading"
                and title_marker.lower() in _block_text(block).lower()):
            found = True
            level = block["level"]
            index += 1
            while index < len(blocks):
                nxt = blocks[index]
                if nxt["kind"] == "heading" and nxt["level"] <= level:
                    break
                index += 1
            continue
        result.append(block)
        index += 1
    if not found:
        raise ValueError(f"Confluence trim marker not found: {title_marker!r}")
    return result


def emit_blocks_to_document(blocks: list[dict]) -> DocumentObject:
    document = Document()
    for block in blocks:
        if block["kind"] == "heading":
            paragraph = document.add_heading("", level=min(block["level"], 4))
            for text, bold, italic in block["runs"]:
                run = paragraph.add_run(text)
                run.bold = bold or None
                run.italic = italic
        elif block["kind"] == "table":
            rows = block["rows"]
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for row_values, row in zip(rows, table.rows):
                for value, cell in zip(row_values, row.cells):
                    cell.text = value
        else:
            style = "List Bullet" if block["kind"] == "li" else None
            paragraph = document.add_paragraph(style=style)
            for text, bold, italic in block["runs"]:
                run = paragraph.add_run(text)
                run.bold = bold or None
                run.italic = italic
    return document


def curate_confluence(source: Path, destination: Path, trims: Iterable[str], expect: str) -> None:
    if not source.is_file():
        raise FileNotFoundError(source)
    blocks = parse_confluence_export(source)

    # Confluence prepends the page name as an extra h1 duplicating the real
    # document title beneath it — the same "remove duplicated boilerplate"
    # rule the native DOCX sources follow.
    if blocks and blocks[0]["kind"] == "heading" and "Template" in _block_text(blocks[0]):
        blocks = blocks[1:]

    for marker in trims:
        blocks = drop_block_section(blocks, marker)

    joined = " ".join(_block_text(block) for block in blocks)
    if expect not in joined:
        raise ValueError(f"{source.name}: expected content marker missing after trim: {expect!r}")

    document = emit_blocks_to_document(blocks)
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    with tempfile.TemporaryDirectory(prefix="handbook-docx-") as temp_dir:
        intermediate = Path(temp_dir) / destination.name
        document.save(intermediate)
        sanitized_package(intermediate, destination)
    Document(destination)


# (source, destination, sections dropped, content marker that must survive)
CONFLUENCE_CURATIONS: tuple[tuple[str, str, tuple[str, ...], str], ...] = (
    (
        "BA/BRD+Template++-+Best+Practice+Version.doc",
        "ba/business-requirements-document.docx",
        ("Lịch sử Thay đổi", "MỤC LỤC"),
        "PHẦN V: YÊU CẦU PHI CHỨC NĂNG",
    ),
    (
        "PO/FRD+Template+-+Best+Practice+Version.doc",
        "po/functional-specification-document.docx",
        ("Lịch sử thay đổi", "Mục lục"),
        "8.3 Đặc tả chuyển đổi dữ liệu",
    ),
    (
        "PO/PRD+Template+-+Best+Practice+Version.doc",
        "po/product-requirements-document.docx",
        ("Lịch sử thay đổi", "Mục lục"),
        "7. Câu hỏi mở",
    ),
)


# ── HTML previews ───────────────────────────────────────────────────────
# The documents are CONFIDENTIAL, so the site's primary interaction is
# *reading in the browser*; the DOCX download is offered from inside the
# reader. Previews are rendered from the already-curated packages (never from
# Template/ working copies), so what the reader shows is byte-derived from
# exactly what the download delivers.

import html as html_escape_module

PREVIEW_STYLE = """
  :root { color-scheme: light dark; }
  body { margin: 0 auto; max-width: 860px; padding: 32px 20px 64px;
         font: 15px/1.65 system-ui, 'Segoe UI', sans-serif;
         background: #ffffff; color: #1c1c1c; }
  h1 { font-size: 24px; line-height: 1.3; margin: 18px 0 6px; }
  h2 { font-size: 19px; margin: 26px 0 8px; }
  h3 { font-size: 16px; margin: 20px 0 6px; }
  h4, h5 { font-size: 14.5px; margin: 16px 0 6px; }
  p { margin: 8px 0; }
  ul { margin: 8px 0 8px 22px; padding: 0; }
  li { margin: 4px 0; }
  table { border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 13.5px; }
  td { border: 1px solid #d0d0d0; padding: 6px 9px; vertical-align: top; }
  tr.doc-preview-thead td { background: #f3f3f3; font-weight: 600; }
  mark { background: #ffe58a; padding: 0 2px; border-radius: 2px; }
  .doc-preview-head { border-bottom: 2px solid #da2128; padding-bottom: 14px; margin-bottom: 18px; }
  .doc-preview-classification { font-weight: 700; color: #b31217; letter-spacing: .04em;
                                text-transform: uppercase; font-size: 12px; margin: 0 0 6px; }
  .doc-preview-note { margin: 0 0 10px; font-size: 13px; color: #555; }
  .doc-preview-standalone-download a { display: inline-block; background: #da2128; color: #fff;
      text-decoration: none; font-weight: 600; font-size: 13.5px; padding: 8px 18px; border-radius: 8px; }
  @media (prefers-color-scheme: dark) {
    body { background: #181818; color: #e6e6e6; }
    td { border-color: #3a3a3a; }
    tr.doc-preview-thead td { background: #262626; }
    mark { background: #7a5d00; color: #fff; }
    .doc-preview-note { color: #a8a8a8; }
  }
"""


def _render_runs(paragraph: Paragraph) -> str:
    parts: list[str] = []
    for run in paragraph.runs:
        text = html_escape_module.escape(run.text)
        if not text.strip():
            if text:
                parts.append(text)
            continue
        if run.font.highlight_color is not None:
            text = f"<mark>{text}</mark>"
        if run.bold:
            text = f"<strong>{text}</strong>"
        if run.italic:
            text = f"<em>{text}</em>"
        parts.append(text)
    rendered = "".join(parts)
    return rendered if rendered.strip() else ""


def _paragraph_tag(paragraph: Paragraph) -> str:
    style_name = paragraph.style.name if paragraph.style is not None else ""
    style_name = style_name or ""
    if style_name == "Title":
        return "h1"
    if style_name.startswith("Heading"):
        try:
            level = int(style_name.rsplit(" ", 1)[-1])
        except ValueError:
            level = 1
        return f"h{min(level + 1, 5)}"
    return "p"


def _is_list_item(paragraph: Paragraph) -> bool:
    properties = paragraph._p.pPr
    if properties is not None and properties.find(qn("w:numPr")) is not None:
        return True
    style_name = paragraph.style.name if paragraph.style is not None else ""
    return (style_name or "").startswith("List")


def render_preview_body(document: DocumentObject) -> str:
    blocks: list[str] = []
    list_items: list[str] = []

    def flush_list() -> None:
        nonlocal list_items
        if list_items:
            items = "".join(f"<li>{item}</li>" for item in list_items)
            blocks.append(f"<ul>{items}</ul>")
            list_items = []

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document)
            rendered = _render_runs(paragraph)
            if not rendered:
                continue
            if _is_list_item(paragraph):
                list_items.append(rendered)
                continue
            flush_list()
            tag = _paragraph_tag(paragraph)
            blocks.append(f"<{tag}>{rendered}</{tag}>")
        elif child.tag == qn("w:tbl"):
            flush_list()
            table = Table(child, document)
            rows_html: list[str] = []
            for row_index, row in enumerate(table.rows):
                cells_html: list[str] = []
                seen_cells: set[int] = set()  # merged cells repeat the same tc
                for cell in row.cells:
                    marker = id(cell._tc)
                    if marker in seen_cells:
                        continue
                    seen_cells.add(marker)
                    cell_parts = []
                    for cell_paragraph in cell.paragraphs:
                        rendered = _render_runs(cell_paragraph)
                        if rendered:
                            cell_parts.append(rendered)
                    cells_html.append(f"<td>{'<br/>'.join(cell_parts) or '&#160;'}</td>")
                row_class = ' class="doc-preview-thead"' if row_index == 0 else ""
                rows_html.append(f"<tr{row_class}>{''.join(cells_html)}</tr>")
            blocks.append(f"<table>{''.join(rows_html)}</table>")
    flush_list()
    return "\n".join(blocks)


def render_preview_page(document: DocumentObject, destination_name: str) -> str:
    body = render_preview_body(document)
    # Titles come from the same run-level renderer as the body: raw body_text()
    # walks mc:AlternateContent Choice AND Fallback branches and repeats banner
    # text, so the naive first-paragraph approach produced tripled titles.
    title_parts: list[str] = []
    for child in document.element.body.iterchildren():
        if child.tag != qn("w:p"):
            continue
        rendered = re.sub(r"<[^>]+>", "", _render_runs(Paragraph(child, document))).strip()
        if rendered:
            title_parts.append(rendered)
        if len(title_parts) == 2:
            break
    title = " — ".join(title_parts) if title_parts else Path(destination_name).stem
    title = html_escape_module.escape(title[:110])
    # previews live in assets/templates/previews/, one level beside the role dirs
    download_href = "../" + destination_name.replace("\\", "/")
    return f"""<!DOCTYPE html>
<html lang="vi">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1"/>
<meta name="referrer" content="no-referrer"/>
<title>{title} — Xem trước</title>
<style>{PREVIEW_STYLE}</style>
</head>
<body>
<article id="doc-preview-content" class="doc-preview">
<header class="doc-preview-head">
<p class="doc-preview-classification">MẬT (CONFIDENTIAL) — Lưu hành nội bộ</p>
<p class="doc-preview-note">Bản xem trước của template DOCX — nội dung sinh trực tiếp từ đúng file sẽ tải về. Đọc tại đây trước; chỉ tải khi thực sự cần dùng.</p>
<p class="doc-preview-standalone-download"><a href="{download_href}" download>Tải bản DOCX</a></p>
</header>
{body}
</article>
</body>
</html>
"""


def _all_destination_names() -> list[str]:
    return [dest for _, dest, _ in CURATIONS] + [dest for _, dest, _, _ in CONFLUENCE_CURATIONS]


def build_previews(output_root: Path) -> list[Path]:
    previews_root = output_root / "previews"
    previews_root.mkdir(parents=True, exist_ok=True)
    written: list[Path] = []
    for destination_name in _all_destination_names():
        docx_path = output_root / destination_name
        if not docx_path.is_file():
            raise FileNotFoundError(f"Curated document missing before preview build: {docx_path}")
        document = Document(docx_path)
        page = render_preview_page(document, destination_name)
        preview_path = previews_root / (Path(destination_name).stem + ".html")
        preview_path.write_text(page, encoding="utf-8", newline="\n")
        written.append(preview_path)
        print(f"[OK] preview {preview_path.relative_to(ROOT)}")
    return written


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output-root",
        type=Path,
        default=DEFAULT_OUTPUT_ROOT,
        help="Destination root (default: assets/templates)",
    )
    parser.add_argument(
        "--clean",
        action="store_true",
        help="Delete the output root before rebuilding the explicit curated set",
    )
    parser.add_argument(
        "--previews-only",
        action="store_true",
        help="Rebuild only the HTML previews from the existing curated DOCX files",
    )
    parser.add_argument(
        "--confluence-only",
        action="store_true",
        help="Curate only the Confluence-export sources (leaves native DOCX outputs untouched), then rebuild previews",
    )
    args = parser.parse_args()
    output_root = args.output_root.resolve()

    if args.previews_only:
        previews = build_previews(output_root)
        print(f"[OK] rendered {len(previews)} HTML previews")
        return

    if args.confluence_only:
        for source_name, destination_name, trims, expect in CONFLUENCE_CURATIONS:
            source = SOURCE_ROOT / source_name
            destination = output_root / destination_name
            curate_confluence(source, destination, trims, expect)
            print(f"[OK] {source.relative_to(ROOT)} -> {destination.relative_to(ROOT)} (Confluence export)")
        previews = build_previews(output_root)
        print(f"[OK] rendered {len(previews)} HTML previews")
        return

    if args.clean and output_root.exists():
        shutil.rmtree(output_root)

    written: list[Path] = []
    for source_name, destination_name, transform in CURATIONS:
        source = SOURCE_ROOT / source_name
        destination = output_root / destination_name
        curate(source, destination, transform)
        written.append(destination)
        print(f"[OK] {source.relative_to(ROOT)} -> {destination.relative_to(ROOT)}")
    for source_name, destination_name, trims, expect in CONFLUENCE_CURATIONS:
        source = SOURCE_ROOT / source_name
        destination = output_root / destination_name
        curate_confluence(source, destination, trims, expect)
        written.append(destination)
        print(f"[OK] {source.relative_to(ROOT)} -> {destination.relative_to(ROOT)} (Confluence export)")
    print(f"[OK] curated {len(written)} publishable DOCX files")
    previews = build_previews(output_root)
    print(f"[OK] rendered {len(previews)} HTML previews")


if __name__ == "__main__":
    main()
